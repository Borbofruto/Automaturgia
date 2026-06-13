"""
docx_generator.py — Geração de documentos Word (.docx) para entregas da Automaturge.

Converte conteúdo Markdown em .docx profissional usando python-docx.
Requer: pip install python-docx markdown

Uso:
    python docx_generator.py --input relatorio.md --output relatorio.docx --title "Título"
    python docx_generator.py --content "# Seção\n\nTexto..." --output saida.docx
"""

import sys
import os
import re
import argparse
from datetime import datetime


def parse_markdown_to_docx(markdown_content: str, output_path: str, title: str = "", author: str = "Automaturge") -> str:
    """Converte Markdown para .docx preservando estrutura de headings, tabelas e listas."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("Instale: pip install python-docx")

    doc = Document()

    # Configuração de página
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # Metadados
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.created = datetime.now()
    doc.core_properties.comments = "Gerado pelo pipeline Automaturge"

    # Título da capa
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Automaturge · Research & Methods for Applied Robotics")
        doc.add_paragraph(datetime.now().strftime("%d/%m/%Y"))
        doc.add_page_break()

    lines = markdown_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        # Tabelas Markdown
        elif "|" in line and i + 1 < len(lines) and "|---" in lines[i + 1]:
            # Coleta cabeçalho
            headers = [h.strip() for h in line.strip("|").split("|")]
            i += 2  # pula linha separadora

            rows = []
            while i < len(lines) and "|" in lines[i]:
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"

            # Cabeçalho
            hdr_row = table.rows[0]
            for j, h in enumerate(headers):
                cell = hdr_row.cells[j]
                cell.text = h
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

            # Dados
            for r_idx, row_data in enumerate(rows):
                row = table.rows[r_idx + 1]
                for j, val in enumerate(row_data):
                    if j < len(row.cells):
                        row.cells[j].text = val

            doc.add_paragraph()
            continue

        # Listas não-ordenadas
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")

        # Listas ordenadas
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")

        # Linha horizontal
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        # Bloco de código
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph("\n".join(code_lines))
            p.style = "No Spacing"
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        # Linha em branco
        elif line.strip() == "":
            pass

        # Parágrafo normal
        else:
            # Processa inline bold/italic
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            doc.add_paragraph(text)

        i += 1

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate(content_markdown: str, output_path: str, title: str = "", author: str = "Automaturge") -> str:
    """
    Ponto de entrada principal.

    Args:
        content_markdown: Conteúdo em Markdown
        output_path:      Caminho .docx de saída
        title:            Título do documento
        author:           Autor (padrão: Automaturge)

    Returns:
        Caminho do arquivo .docx gerado
    """
    return parse_markdown_to_docx(content_markdown, output_path, title, author)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Automaturge DOCX Generator")
    parser.add_argument("--input", help="Arquivo .md de entrada")
    parser.add_argument("--content", help="Conteúdo Markdown direto")
    parser.add_argument("--output", required=True, help="Arquivo .docx de saída")
    parser.add_argument("--title", default="", help="Título do documento")
    parser.add_argument("--author", default="Automaturge", help="Autor")
    args = parser.parse_args()

    if args.input:
        with open(args.input, encoding="utf-8") as f:
            content = f.read()
    elif args.content:
        content = args.content
    else:
        content = sys.stdin.read()

    result = generate(content, args.output, args.title, args.author)
    print(f"DOCX gerado: {result}")
